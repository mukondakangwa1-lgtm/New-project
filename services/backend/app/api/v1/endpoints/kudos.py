"""
Digital Campus - KUDOS AI Assistant
Document learning, web learning, retrieval-based chat, superadmin controls.
"""
import io
import json
import re
from collections import OrderedDict
from datetime import datetime, timezone
from typing import Optional

import httpx
from bs4 import BeautifulSoup
from fastapi import APIRouter, Depends, HTTPException, UploadFile, File, Form
from sqlalchemy import func, or_
from sqlalchemy.orm import Session, joinedload

from app.core.database import get_db
from app.core.deps import get_current_user, require_admin
from app.core.kudos_guardian import self_improver
from app.models import (
    KudosChunk, KudosConversation, KudosDocument, KudosMessage, KudosWebKnowledge, User,
)
from app.schemas import (
    KudosAskRequest, KudosAskResponse, KudosConversationResponse, KudosDocumentResponse,
    KudosDocumentUpdate, KudosMessageResponse, KudosStats, KudosWebKnowledgeResponse, KudosWebLearn,
)

router = APIRouter()

CHUNK_SIZE = 500
CHUNK_OVERLAP = 50
STOP_WORDS = set(
    "the a an and or but in on at to for of is it that this with from by as are was were be been "
    "being have has had do does did will would shall should may might can could i me my we our you "
    "your he she they them their its not no nor so if than too very just about above after again all "
    "any because before between both each few more most other some such then there these through "
    "under until when where which while who whom why how".split()
)


def extract_text_from_file(content: bytes, filename: str) -> str:
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext in ("txt", "md", "csv", "json", "py", "js", "html", "css"):
        return content.decode("utf-8", errors="ignore")
    if ext == "pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(io.BytesIO(content))
            return "".join(page.extract_text() or "" for page in reader.pages)
        except Exception:
            return content.decode("utf-8", errors="ignore")
    if ext in ("docx", "doc"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception:
            return content.decode("utf-8", errors="ignore")
    return content.decode("utf-8", errors="ignore")


def chunk_text(text: str, chunk_size: int = CHUNK_SIZE, overlap: int = CHUNK_OVERLAP) -> list[str]:
    words = text.split()
    if len(words) <= chunk_size:
        return [text.strip()] if text.strip() else []
    chunks = []
    start = 0
    while start < len(words):
        end = start + chunk_size
        chunk = " ".join(words[start:end])
        if chunk.strip():
            chunks.append(chunk.strip())
        start += chunk_size - overlap
    return chunks


def extract_keywords(text: str, max_keywords: int = 20) -> str:
    words = re.findall(r"[a-zA-Z]{3,}", text.lower())
    freq: dict[str, int] = {}
    for w in words:
        if w not in STOP_WORDS:
            freq[w] = freq.get(w, 0) + 1
    return ",".join(w for w, _ in sorted(freq.items(), key=lambda x: -x[1])[:max_keywords])


def simple_summarize(text: str, max_sentences: int = 5) -> str:
    sentences = re.split(r"[.!?]+", text)
    sentences = [s.strip() for s in sentences if len(s.strip()) > 20]
    return ". ".join(sentences[:max_sentences]) + "." if sentences else text[:500]


def search_chunks(db: Session, query: str, limit: int = 5) -> list[dict]:
    query_words = set(re.findall(r"[a-zA-Z]{3,}", query.lower())) - STOP_WORDS
    if not query_words:
        return []
    try:
        chunks = db.query(KudosChunk).join(KudosDocument).filter(
            KudosDocument.is_approved == True, KudosDocument.is_active == True
        ).all()
    except Exception:
        return []

    scored = []
    for chunk in chunks:
        content_lower = chunk.content.lower()
        keywords = set(chunk.keywords.split(",")) if chunk.keywords else set()
        score = sum(3 if w in keywords else 1 for w in query_words if w in content_lower)
        if score > 0:
            scored.append({"chunk_id": chunk.id, "document_id": chunk.document_id, "content": chunk.content[:500], "score": score})

    try:
        web_items = db.query(KudosWebKnowledge).filter(
            KudosWebKnowledge.is_approved == True, KudosWebKnowledge.is_active == True
        ).all()
        for item in web_items:
            content_lower = (item.content or "").lower()
            score = sum(1 for w in query_words if w in content_lower)
            if score > 0:
                scored.append({"chunk_id": None, "document_id": None, "web_id": item.id, "title": item.title, "content": (item.summary or item.content[:500])[:500], "score": score})
    except Exception:
        pass

    scored.sort(key=lambda x: -x["score"])
    return scored[:limit]


def generate_answer(query: str, sources: list[dict]) -> str:
    """Simple fallback answer — never used if conversation engine works."""
    if not sources:
        return f"I don't have information about \"{query}\" yet. Try uploading a document or teaching me a web page about it."
    # Extract relevant content
    import re
    query_words = set(re.findall(r"[a-zA-Z]{3,}", query.lower())) - STOP_WORDS
    best_content = ""
    best_score = 0
    for src in sources:
        content = src.get("content", "")
        score = sum(1 for w in query_words if w in content.lower())
        if score > best_score:
            best_score = score
            best_content = content
    # Extract relevant sentences
    sentences = re.split(r'[.!?\n]+', best_content)
    relevant = [s.strip() for s in sentences if len(s.strip()) > 20 and any(w in s.lower() for w in query_words)][:3]
    if relevant:
        return ". ".join(relevant) + "."
    return best_content[:400]


# ──────────────────────────────────────────────
# DOCUMENT ENDPOINTS
# ──────────────────────────────────────────────


@router.post("/documents/upload", response_model=KudosDocumentResponse, status_code=201)
async def upload_document(
    title: str = Form(...), tags: str = Form(""), file: UploadFile = File(...),
    db: Session = Depends(get_db), current_user: User = Depends(get_current_user),
):
    content_bytes = await file.read()
    if len(content_bytes) > 10 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="File too large (max 10MB)")
    text = extract_text_from_file(content_bytes, file.filename or "unknown.txt")
    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")
    doc = KudosDocument(
        uploaded_by=current_user.id, title=title, filename=file.filename or "unknown",
        file_type=file.filename.rsplit(".", 1)[-1].lower() if "." in (file.filename or "") else "",
        content=text, summary=simple_summarize(text), tags=tags, is_approved=current_user.is_admin,
    )
    db.add(doc)
    db.flush()
    for i, chunk_content in enumerate(chunk_text(text)):
        db.add(KudosChunk(document_id=doc.id, chunk_index=i, content=chunk_content, word_count=len(chunk_content.split()), keywords=extract_keywords(chunk_content)))
    doc.chunk_count = len(chunk_text(text))
    db.commit()
    db.refresh(doc)
    return doc


@router.get("/documents", response_model=list[KudosDocumentResponse])
def list_documents(show_all: bool = False, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    q = db.query(KudosDocument)
    if not current_user.is_admin:
        q = q.filter(KudosDocument.is_approved == True, KudosDocument.is_active == True)
    elif not show_all:
        q = q.filter(KudosDocument.is_active == True)
    return q.order_by(KudosDocument.created_at.desc()).all()


@router.patch("/documents/{doc_id}", response_model=KudosDocumentResponse)
def update_document(doc_id: int, body: KudosDocumentUpdate, db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    doc = db.query(KudosDocument).filter(KudosDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    for field, value in body.model_dump(exclude_unset=True).items():
        setattr(doc, field, value)
    db.commit()
    db.refresh(doc)
    return doc


@router.delete("/documents/{doc_id}", status_code=204)
def delete_document(doc_id: int, db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    doc = db.query(KudosDocument).filter(KudosDocument.id == doc_id).first()
    if not doc:
        raise HTTPException(status_code=404, detail="Document not found")
    db.delete(doc)
    db.commit()


# ──────────────────────────────────────────────
# WEB LEARNING ENDPOINTS
# ──────────────────────────────────────────────


@router.post("/learn/web", response_model=KudosWebKnowledgeResponse, status_code=201)
async def learn_web_page(body: KudosWebLearn, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    if not body.url.startswith(("http://", "https://")):
        raise HTTPException(status_code=400, detail="Invalid URL")
    try:
        async with httpx.AsyncClient(timeout=10, follow_redirects=True) as client:
            response = await client.get(body.url)
            response.raise_for_status()
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to fetch URL: {e}")
    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "nav", "footer", "header"]):
        tag.decompose()
    title = body.title or (soup.title.string if soup.title else body.url)
    text = soup.get_text(separator="\n", strip=True)
    if len(text) < 50:
        raise HTTPException(status_code=400, detail="Page has too little text content")
    knowledge = KudosWebKnowledge(url=body.url, title=title[:255], content=text, summary=simple_summarize(text), is_approved=current_user.is_admin, learned_by=current_user.id)
    db.add(knowledge)
    db.commit()
    db.refresh(knowledge)
    return knowledge


@router.get("/learn/web", response_model=list[KudosWebKnowledgeResponse])
def list_web_knowledge(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    q = db.query(KudosWebKnowledge)
    if not current_user.is_admin:
        q = q.filter(KudosWebKnowledge.is_approved == True, KudosWebKnowledge.is_active == True)
    return q.order_by(KudosWebKnowledge.created_at.desc()).all()


@router.patch("/learn/web/{item_id}", response_model=KudosWebKnowledgeResponse)
def update_web_knowledge(item_id: int, body: KudosDocumentUpdate, db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    item = db.query(KudosWebKnowledge).filter(KudosWebKnowledge.id == item_id).first()
    if not item:
        raise HTTPException(status_code=404, detail="Web knowledge not found")
    for field, value in body.model_dump(exclude_unset=True).items():
        if hasattr(item, field):
            setattr(item, field, value)
    db.commit()
    db.refresh(item)
    return item


# ──────────────────────────────────────────────
# CHAT WITH KUDOS — ROBUST ERROR HANDLING
# ──────────────────────────────────────────────


@router.post("/ask", response_model=KudosAskResponse)
async def ask_kudos(body: KudosAskRequest, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    """Ask KUDOS — always returns an answer, never crashes."""
    try:
        # Get or create conversation
        conv = None
        if body.conversation_id:
            try:
                conv = db.query(KudosConversation).filter(
                    KudosConversation.id == body.conversation_id,
                    KudosConversation.user_id == current_user.id,
                ).first()
            except Exception:
                conv = None
        if not conv:
            conv = KudosConversation(user_id=current_user.id, title=body.question[:100])
            db.add(conv)
            db.flush()

        # Save user message
        db.add(KudosMessage(conversation_id=conv.id, role="user", content=body.question))

        # Search knowledge base
        sources = []
        try:
            sources = search_chunks(db, body.question)
        except Exception:
            pass

        # Build knowledge context from sources
        knowledge_context = ""
        if sources:
            knowledge_context = "\n".join(s.get("content", "")[:300] for s in sources[:3])

        # Try LLM first (human-like response)
        answer = ""
        try:
            from app.core.llm_engine import get_llm_response
            conv_history = []
            try:
                conv_history = db.query(KudosMessage).filter(
                    KudosMessage.conversation_id == conv.id
                ).order_by(KudosMessage.created_at.desc()).limit(5).all()
                conv_history = [{"role": m.role, "content": m.content} for m in conv_history]
            except Exception:
                pass

            llm_answer = await get_llm_response(
                question=body.question,
                knowledge_context=knowledge_context,
                conversation_history=conv_history,
                user_name=current_user.full_name.split()[0] if current_user.full_name else "",
            )
            if llm_answer and len(llm_answer) > 10:
                answer = llm_answer
        except Exception:
            pass

        # Fallback to internal engine
        if not answer or len(answer) < 10:
            try:
                from app.core.conversation_engine import generate_human_response
                answer = generate_human_response(
                    query=body.question, sources=sources, conv_id=conv.id,
                    user_name=current_user.full_name.split()[0] if current_user.full_name else None,
                )
            except Exception:
                answer = generate_answer(body.question, sources)

        if not answer or len(answer) < 10:
            answer = generate_answer(body.question, sources)

        # Self-improvement logging
        try:
            self_improver.log_question(current_user.id, body.question, had_sources=bool(sources))
        except Exception:
            pass

        # Save KUDOS response
        try:
            db.add(KudosMessage(
                conversation_id=conv.id, role="kudos", content=answer,
                sources=json.dumps(sources[:3]) if sources else "[]",
            ))
            db.commit()
        except Exception:
            db.rollback()

        return KudosAskResponse(
            answer=answer,
            sources=[
                {"document_id": s.get("document_id"), "web_id": s.get("web_id"), "title": s.get("title", ""), "preview": s.get("content", "")[:200]}
                for s in (sources[:3] if sources else [])
            ],
            conversation_id=conv.id,
        )

    except Exception as e:
        try:
            db.rollback()
        except Exception:
            pass
        return KudosAskResponse(
            answer=f"I had trouble processing that. Please try again. ({str(e)[:100]})",
            sources=[], conversation_id=body.conversation_id or 0,
        )


@router.get("/conversations", response_model=list[KudosConversationResponse])
def list_conversations(db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    return db.query(KudosConversation).filter(KudosConversation.user_id == current_user.id).order_by(KudosConversation.created_at.desc()).all()


@router.get("/conversations/{conv_id}/messages", response_model=list[KudosMessageResponse])
def get_conversation_messages(conv_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    conv = db.query(KudosConversation).filter(KudosConversation.id == conv_id, KudosConversation.user_id == current_user.id).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    return db.query(KudosMessage).filter(KudosMessage.conversation_id == conv_id).order_by(KudosMessage.created_at).all()


@router.delete("/conversations/{conv_id}", status_code=204)
def delete_conversation(conv_id: int, db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    conv = db.query(KudosConversation).filter(KudosConversation.id == conv_id, KudosConversation.user_id == current_user.id).first()
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
    db.delete(conv)
    db.commit()


# ──────────────────────────────────────────────
# SUPERADMIN CONTROLS
# ──────────────────────────────────────────────


@router.get("/admin/stats", response_model=KudosStats)
def kudos_stats(db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    return KudosStats(
        total_documents=db.query(KudosDocument).count(),
        approved_documents=db.query(KudosDocument).filter(KudosDocument.is_approved == True).count(),
        total_chunks=db.query(KudosChunk).count(),
        total_web_knowledge=db.query(KudosWebKnowledge).count(),
        total_conversations=db.query(KudosConversation).count(),
        total_messages=db.query(KudosMessage).count(),
    )


@router.post("/admin/approve-all-documents")
def approve_all_documents(db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    count = db.query(KudosDocument).filter(KudosDocument.is_approved == False).update({"is_approved": True})
    db.commit()
    return {"approved": count}


@router.post("/admin/approve-all-web")
def approve_all_web(db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    count = db.query(KudosWebKnowledge).filter(KudosWebKnowledge.is_approved == False).update({"is_approved": True})
    db.commit()
    return {"approved": count}


@router.post("/admin/pending")
def list_pending(db: Session = Depends(get_db), admin: User = Depends(require_admin)):
    docs = db.query(KudosDocument).filter(KudosDocument.is_approved == False).all()
    web = db.query(KudosWebKnowledge).filter(KudosWebKnowledge.is_approved == False).all()
    return {
        "pending_documents": [{"id": d.id, "title": d.title, "uploaded_by": d.uploaded_by, "chunks": d.chunk_count} for d in docs],
        "pending_web": [{"id": w.id, "url": w.url, "title": w.title, "learned_by": w.learned_by} for w in web],
    }
